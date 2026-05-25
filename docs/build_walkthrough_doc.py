"""Build a humanized Word-doc walkthrough of how VitalScan works under the hood.

Output:  docs/VitalScan_Technical_Walkthrough.docx
Run:     backend/venv312/bin/python docs/build_walkthrough_doc.py
"""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "VitalScan_Technical_Walkthrough.docx"

NAVY = RGBColor(0x0F, 0x1F, 0x44)
ACCENT = RGBColor(0x1F, 0x6F, 0xEB)
MUTED = RGBColor(0x55, 0x66, 0x7A)
GOOD = RGBColor(0x1E, 0x88, 0x4F)
INK = RGBColor(0x1A, 0x1A, 0x1A)


# ---------- helpers ----------
_bookmark_counter = [0]
_pending_bookmark = [None]   # set before next add_heading() to attach bookmark


def _next_bookmark_id() -> int:
    _bookmark_counter[0] += 1
    return _bookmark_counter[0]


def _attach_bookmark(paragraph, name: str):
    """Wrap the paragraph's runs with a Word bookmark anchor."""
    bm_id = str(_next_bookmark_id())
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bm_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _rgb_to_hex(c: RGBColor) -> str:
    return "{:02X}{:02X}{:02X}".format(*c)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *,
                           bold=False, size=11, color=NAVY,
                           font_name="Calibri"):
    """Add a clickable internal hyperlink that jumps to a bookmark anchor."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    col_el = OxmlElement("w:color")
    col_el.set(qn("w:val"), _rgb_to_hex(color))
    rPr.append(col_el)
    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1, color=NAVY, *, bookmark=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    # Resolve a bookmark name from: explicit param > pending queue > TOC title lookup.
    # The TOC lookup strips the "N." or "N.M" number prefix because add_heading() calls
    # in the script use "1. What this document is" while TOC_ENTRIES separates them.
    bm = bookmark or _pending_bookmark[0]
    if bm is None:
        # try lookup: strip "1." / "2.1" / "5.11" prefix, then match
        parts = text.split(" ", 1)
        if (len(parts) > 1
                and parts[0].rstrip(".").replace(".", "").isdigit()):
            title_only = parts[1].strip()  # double-space tolerated
        else:
            title_only = text.strip()
        bm = TITLE_TO_BOOKMARK.get(title_only) or TITLE_TO_BOOKMARK.get(text)
    if bm is not None:
        _attach_bookmark(h, bm)
        _pending_bookmark[0] = None
    return h


def add_para(doc, text, *, size=11, italic=False, bold=False, color=INK,
             align=None, space_after=8):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        # nested bullet via numbering level
        p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    return p


def add_callout(doc, label, body, *, label_color=ACCENT):
    """Indented italic callout — used for analogies and aside notes."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(body)
    r2.italic = True
    r2.font.color.rgb = INK
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    # subtle gray background-ish color via lighter text
    r.font.color.rgb = RGBColor(0x33, 0x44, 0x5C)


def add_table(doc, rows, header_row=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            if i == 0 and header_row:
                r.bold = True
                r.font.color.rgb = NAVY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


# ---------- TOC entries: (number, title, level, bookmark slug) ----------
TOC_ENTRIES = [
    ("1.",  "What this document is",                                             1, "sec_1"),
    ("2.",  "The stream — how the video gets from your camera to our backend",   1, "sec_2"),
    ("2.1", "The browser asks for camera permission",                            2, "sec_2_1"),
    ("2.2", "MediaRecorder turns the live stream into a file",                   2, "sec_2_2"),
    ("2.3", "The upload — multipart form-data",                                  2, "sec_2_3"),
    ("2.4", "The network path — Cloudflare, the tunnel, the homelab",            2, "sec_2_4"),
    ("3.",  "Video file formats — what each extension actually is",              1, "sec_3"),
    ("4.",  "How the video gets parsed — frame by frame",                        1, "sec_4"),
    ("4.1", "What a frame actually is",                                          2, "sec_4_1"),
    ("4.2", "The for-loop that does the work",                                   2, "sec_4_2"),
    ("5.",  "The technical terms decoded — a humanized glossary",                1, "sec_5"),
    ("5.1", "OpenCV",                                                            2, "sec_5_1"),
    ("5.2", "MediaPipe",                                                         2, "sec_5_2"),
    ("5.3", "FaceMesh — the specific MediaPipe model we use",                    2, "sec_5_3"),
    ("5.4", "Polygon and ROI",                                                   2, "sec_5_4"),
    ("5.5", "Mean RGB",                                                          2, "sec_5_5"),
    ("5.6", "POS algorithm",                                                     2, "sec_5_6"),
    ("5.7", "Bandpass filter",                                                   2, "sec_5_7"),
    ("5.8", "FFT — Fast Fourier Transform",                                      2, "sec_5_8"),
    ("5.9", "Peak detection and IBI",                                            2, "sec_5_9"),
    ("5.10","SDNN — what HRV actually means",                                    2, "sec_5_10"),
    ("5.11","LF/HF ratio and the stress index",                                  2, "sec_5_11"),
    ("6.",  "The math, all in one place",                                        1, "sec_6"),
    ("6.1", "Mean RGB per frame",                                                2, "sec_6_1"),
    ("6.2", "POS projection",                                                    2, "sec_6_2"),
    ("6.3", "Bandpass filter",                                                   2, "sec_6_3"),
    ("6.4", "Heart rate from FFT",                                               2, "sec_6_4"),
    ("6.5", "SDNN",                                                              2, "sec_6_5"),
    ("6.6", "Stress index",                                                      2, "sec_6_6"),
    ("7.",  "A concrete end-to-end example",                                     1, "sec_7"),
    ("8.",  "Why each design choice was made",                                   1, "sec_8"),
    ("9.",  "Confidence per biomarker — how much should you trust each number?", 1, "sec_9"),
    ("9.1", "Heart rate confidence — FFT peak prominence",                       2, "sec_9_1"),
    ("9.2", "HRV and stress confidence — beat count + IBI plausibility",         2, "sec_9_2"),
    ("9.3", "Blood pressure confidence — a low floor by design",                 2, "sec_9_3"),
    ("9.4", "How Group 3 should consume confidence",                             2, "sec_9_4"),
    ("10.", "Common questions, answered",                                        1, "sec_10"),
    ("10.1","What is RMSE, and why use MAE and RMSE together?",                  2, "sec_10_1"),
    ("10.2","Train / test / evaluation split",                                   2, "sec_10_2"),
    ("10.3","Why a mathematical approach instead of training a model?",          2, "sec_10_3"),
    ("10.4","How does the dataset map videos to biometric values?",              2, "sec_10_4"),
    ("10.5","Why MediaPipe FaceMesh instead of dlib or MTCNN?",                  2, "sec_10_5"),
    ("11.", "The whole project, in one paragraph",                               1, "sec_11"),
]

# heading-title → bookmark slug — when add_heading() sees a matching title,
# it auto-attaches the bookmark.
TITLE_TO_BOOKMARK = {title: slug for _, title, _, slug in TOC_ENTRIES}


def add_toc(doc):
    """Render the TOC with every entry as a clickable internal hyperlink."""
    for num, title, level, anchor in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if level == 2 else 4)
        p.paragraph_format.left_indent = Cm(0.0 if level == 1 else 0.8)
        # Number column — also part of the hyperlink so the whole row is clickable
        add_internal_hyperlink(
            p, anchor, f"{num:5s}  ",
            bold=(level == 1),
            size=11 if level == 1 else 10,
            color=ACCENT if level == 1 else MUTED,
        )
        # Title column
        add_internal_hyperlink(
            p, anchor, title,
            bold=(level == 1),
            size=11 if level == 1 else 10,
            color=NAVY if level == 1 else INK,
        )


# ============================================================
# BUILD THE DOCUMENT
# ============================================================
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

# --- Cover ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("VitalScan — Under the Hood")
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = NAVY
tr.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("How the video stream is captured, parsed, processed,\nand turned into biomarker JSON — in plain English.")
sr.font.size = Pt(13)
sr.font.color.rgb = MUTED
sr.font.name = "Calibri"
sr.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("AIT 500  ·  Group 1  ·  rPPG Signal Extraction")
mr.font.size = Pt(11)
mr.font.color.rgb = ACCENT
mr.font.name = "Calibri"

doc.add_paragraph()
doc.add_paragraph()

# --- Table of Contents ---
add_heading(doc, "Table of Contents", level=1)
add_toc(doc)
doc.add_page_break()

# ============================================================
# 1. WHAT THIS DOCUMENT IS
# ============================================================
add_heading(doc, "1. What this document is", level=1)
add_para(doc,
    "This is the plain-English version of how the VitalScan project actually "
    "works. If you opened the codebase and got hit with terms like \"FaceMesh,\" "
    "\"polygon ROI,\" \"POS projection,\" \"FFT,\" and \"mean RGB,\" this is the "
    "document that explains what every one of those things means — and why we "
    "use them — without assuming you have a computer-vision background.")
add_para(doc,
    "The goal isn't to teach you signal processing from scratch. The goal is "
    "to give you enough mental model that when someone asks \"how does your "
    "project work,\" you can walk them through it in a way that makes sense to "
    "a human being.")

add_callout(doc, "Anchor sentence",
    "We take a 30-second video of someone's face, and we return their heart rate, "
    "heart-rate variability, and a stress score. The whole project is a chain of "
    "steps that turn pixels into vital signs.")

# ============================================================
# 2. THE STREAM
# ============================================================
add_heading(doc, "2. The stream — how the video gets from your camera to our backend", level=1)

add_para(doc,
    "When you click \"Scan live\" on vitalscan.bkre8tive.com, a chain of things "
    "happens in order. Let's walk through it.")

add_heading(doc, "2.1  The browser asks for camera permission", level=2)
add_para(doc,
    "Modern browsers have a built-in API called getUserMedia. When the React app "
    "calls this function, the browser pops up the standard \"vitalscan.bkre8tive.com "
    "wants to use your camera\" prompt. You click Allow, and the browser hands the "
    "React app a live stream of video frames coming off your webcam.")
add_callout(doc, "In plain words",
    "The browser is the middleman between your camera and our code. We never "
    "touch your camera directly — we just ask the browser \"please give us frames,\" "
    "and you have to say yes first.")

add_heading(doc, "2.2  MediaRecorder turns the live stream into a file", level=2)
add_para(doc,
    "A live video stream isn't a file — it's an open faucet of data. To send it "
    "anywhere, we need to record it into a file format. That's what the browser's "
    "MediaRecorder API does: it watches the live stream for 30 seconds and packages "
    "the frames into a single WebM video file in memory.")
add_para(doc,
    "WebM is just one of many container formats — it's the one browsers can produce "
    "natively without needing a download. A 30-second WebM is typically 5–15 MB, "
    "which is small enough to upload over any internet connection in a couple seconds.")

add_heading(doc, "2.3  The upload — multipart form-data", level=2)
add_para(doc,
    "Once we have a video file, we POST it to our backend at /api/scan. The HTTP "
    "request uses something called multipart/form-data — a format designed exactly "
    "for uploading files alongside other form fields. In our case we send two "
    "things in one request:")
add_bullet(doc, "video — the actual WebM file (the big part)")
add_bullet(doc, "person — your name as a text field (the small part, for tagging the scan)")
add_para(doc,
    "We use XMLHttpRequest instead of the more modern fetch API because XHR gives "
    "us upload progress events — we want the green progress bar to show real "
    "bytes-uploaded, not a generic spinner.")

add_heading(doc, "2.4  The network path — Cloudflare, the tunnel, the homelab", level=2)
add_para(doc,
    "The request doesn't go directly to my server. It takes a four-hop path:")
add_bullet(doc, "Browser → Cloudflare edge (TLS terminates here — the HTTPS certificate is Cloudflare's, not ours)")
add_bullet(doc, "Cloudflare edge → Cloudflare Tunnel (an outbound-only secure tunnel that my server initiated)")
add_bullet(doc, "Tunnel → my homelab Docker container (no open ports, no public IP)")
add_bullet(doc, "Container → FastAPI backend on port 8000")
add_callout(doc, "Why this matters",
    "My home firewall is completely closed to incoming traffic. The tunnel "
    "is a phone call my server makes outbound to Cloudflare; Cloudflare can talk "
    "back on that same call, but nobody on the public internet can initiate a "
    "connection to my homelab. This is a much safer pattern than port-forwarding.")

doc.add_page_break()

# ============================================================
# 3. FILE FORMATS
# ============================================================
add_heading(doc, "3. Video file formats — what each extension actually is", level=1)

add_para(doc,
    "Our API accepts five video extensions. Each one is a \"container format\" — "
    "a wrapper that holds compressed video frames, audio (which we ignore), and "
    "metadata. The container is just packaging; the actual video frames inside "
    "are compressed with a separate \"codec.\"")

add_callout(doc, "Container vs codec",
    "Think of it like a cardboard box (the container) holding shrink-wrapped "
    "items (the codec-compressed frames). The .mp4 box is common. The .webm box is "
    "what browsers naturally produce. Inside either box, the actual video could be "
    "encoded with H.264, VP8, VP9, AV1, or many others. We don't care about the "
    "box — we care that we can open it and read the frames out.")

add_table(doc, [
    ["Extension", "Container", "Typical use", "Notes"],
    [".webm", "WebM", "Browser MediaRecorder output", "What our live webcam scans produce"],
    [".mp4", "MPEG-4", "Phones, cameras, screen recordings", "Most common upload format"],
    [".mov", "QuickTime", "Apple devices (iPhone, Mac)", "Same as MP4 internally, Apple wrapper"],
    [".avi", "AVI", "Older Windows tools, dataset videos", "The UBFC-rPPG dataset uses this"],
    [".mkv", "Matroska", "Power-user video tools", "Accepted but rare from end users"],
])

add_para(doc,
    "OpenCV — the library we use to read frames out of any of these — handles all "
    "five transparently. Whatever extension comes in, OpenCV opens it, gives us "
    "one frame at a time as a numpy array of pixel values, and we don't have to "
    "write a single line of format-specific code.")

# ============================================================
# 4. HOW THE VIDEO GETS PARSED
# ============================================================
add_heading(doc, "4. How the video gets parsed — frame by frame", level=1)

add_para(doc,
    "Once the file arrives on the backend, FastAPI saves it to a temporary file on "
    "disk. Then we hand the file path to OpenCV and start reading frames one at a time.")

add_heading(doc, "4.1  What a frame actually is", level=2)
add_para(doc,
    "A video is just a sequence of still images shown rapidly enough that your "
    "eye perceives motion. A 30-second video recorded at 30 frames per second is "
    "900 individual images stacked back-to-back.")
add_para(doc,
    "Each image, in computer memory, is a 3D grid of numbers: height × width × 3. "
    "The 3 is for the red, green, and blue channels. Every pixel is three integers "
    "between 0 and 255 that tell you how much red, green, and blue light bounced "
    "off that point on whatever was in front of the camera.")

add_callout(doc, "Think of it this way",
    "If your video is 640 by 480 pixels at 30 fps for 30 seconds, that's "
    "640 × 480 × 3 × 900 = roughly 830 million numbers per video. That's why "
    "video files have to be compressed — the raw pixel data is enormous.")

add_heading(doc, "4.2  The for-loop that does the work", level=2)
add_para(doc,
    "Once OpenCV opens the file, we run a simple loop:")

add_code(doc,
    "cap = cv2.VideoCapture(video_path)\n"
    "while True:\n"
    "    ret, frame = cap.read()    # read one frame\n"
    "    if not ret:                # ret is False when video is done\n"
    "        break\n"
    "    process_one_frame(frame)   # find face, sample skin colors\n"
    "cap.release()")

add_para(doc,
    "Inside process_one_frame, we do four things — explained in detail in the "
    "next sections:")
add_bullet(doc, "Find the face")
add_bullet(doc, "Identify the forehead and cheeks")
add_bullet(doc, "Average the color values of just those skin regions")
add_bullet(doc, "Save that averaged color as one row in a long table")
add_para(doc,
    "After 900 frames, that long table is 900 rows of (red, green, blue) numbers — "
    "this is the input to the signal-processing stage.")

doc.add_page_break()

# ============================================================
# 5. THE TECHNICAL TERMS DECODED
# ============================================================
add_heading(doc, "5. The technical terms decoded — a humanized glossary", level=1)
add_para(doc,
    "Every weird word in our codebase, defined in plain English.")

add_heading(doc, "5.1  OpenCV", level=2)
add_para(doc,
    "OpenCV is the swiss-army-knife library for reading images and videos in Python. "
    "It was originally written in C++ in the late 1990s for industrial vision tasks "
    "(robotics, factory inspection) and has been the standard ever since. We use it "
    "for one thing only: opening video files and giving us frames as numpy arrays.")
add_callout(doc, "Real-world equivalent",
    "OpenCV is to video what Pillow is to JPEGs. It's the layer that abstracts "
    "away the file format so we can think in terms of pixel arrays.")

add_heading(doc, "5.2  MediaPipe", level=2)
add_para(doc,
    "MediaPipe is Google's open-source library for real-time computer-vision tasks. "
    "It ships pre-trained models for face detection, hand tracking, pose estimation, "
    "and a few others. Crucially, MediaPipe runs on the CPU at video speed — you "
    "don't need a graphics card. That matters because our backend runs in a Docker "
    "container on a homelab without a GPU.")

add_heading(doc, "5.3  FaceMesh — the specific MediaPipe model we use", level=2)
add_para(doc,
    "FaceMesh is one specific model inside MediaPipe. Give it an image with a "
    "face in it, and it returns 468 numbered dots positioned on specific anatomical "
    "landmarks: tip of the nose, corner of the mouth, edge of the left eye, point "
    "on the forehead, and so on.")
add_para(doc,
    "The 468 landmarks are always in the same order. Landmark 10 is always the "
    "same spot on the forehead. Landmark 117 is always the same spot on the left "
    "cheek. This consistency is what lets us reliably extract \"the forehead\" "
    "from every single frame.")
add_callout(doc, "Why 468 specifically",
    "Google chose 468 because it's enough to describe the full 3D shape of a "
    "face — eyebrows, lips, jawline, ear position, all of it — while still being "
    "fast enough to compute 30 times per second on a laptop CPU.")

add_heading(doc, "5.4  Polygon and ROI", level=2)
add_para(doc,
    "A polygon is just a multi-sided shape — like a triangle, a pentagon, or any "
    "freeform many-sided outline. In our code, we build polygons by connecting "
    "specific FaceMesh landmarks. For example, the \"forehead polygon\" is the "
    "outline you get by connecting landmarks 10, 67, 69, 109, 151, 297, 299, 333, "
    "and 337 in order. That outline encloses the forehead skin patch.")
add_para(doc,
    "ROI stands for \"region of interest.\" It's just a fancy term for \"the "
    "specific area we care about right now.\" We have three ROIs: the forehead, "
    "the left cheek, and the right cheek. We ignore everything else in the frame — "
    "the eyes, the mouth, the background, the hair — because those areas don't "
    "give us a clean pulse signal.")

add_heading(doc, "5.5  Mean RGB", level=2)
add_para(doc,
    "Each pixel inside our ROI has three numbers: a red value, a green value, "
    "and a blue value. \"Mean RGB\" means we take all those pixels, average their "
    "red values together to get one red number, average their greens to get one "
    "green number, and average their blues to get one blue number.")
add_para(doc,
    "So for every single frame, we collapse \"all the pixels in the forehead and "
    "both cheeks\" into one (R, G, B) triple. That's the whole point of the ROI "
    "step — reducing thousands of pixels to three numbers per frame.")
add_callout(doc, "Why averaging helps",
    "The pulse signal — the actual color change caused by your heartbeat — is "
    "tiny, around 0.1% to 1% of the total color. Single pixels are too noisy to "
    "see it. But when you average a thousand skin pixels together, the random "
    "pixel-level noise cancels out, and the tiny consistent pulse signal survives.")

add_heading(doc, "5.6  POS algorithm", level=2)
add_para(doc,
    "POS stands for \"Plane-Orthogonal-to-Skin.\" It's a clever 2017 algorithm "
    "from Wang and colleagues at the University of Eindhoven. The problem it "
    "solves is this: when you have a time-series of RGB averages from a face "
    "video, the colors fluctuate for many reasons.")
add_bullet(doc, "Ambient lighting changes (someone walks by, a cloud moves)")
add_bullet(doc, "Head movement (camera angle relative to skin changes)")
add_bullet(doc, "Reflection on the skin's surface (forehead shine)")
add_bullet(doc, "Actual blood pulse (the thing we want)")
add_para(doc,
    "All those things move the RGB values. POS uses a mathematical trick — "
    "projecting the three-dimensional RGB data onto a two-dimensional plane that "
    "is perpendicular to the natural skin-tone color — to subtract out the noise "
    "categories while keeping the pulse.")
add_callout(doc, "Plain analogy",
    "Imagine standing in a noisy room trying to hear one specific person's "
    "voice. POS is like a special microphone that points exactly perpendicular "
    "to the background hum — the hum cancels out, the voice you want comes through.")

add_heading(doc, "5.7  Bandpass filter", level=2)
add_para(doc,
    "A bandpass filter is an electronics concept: it lets through frequencies "
    "in a specific band and blocks everything outside. In our case, the human "
    "heart beats between 42 and 240 times per minute (the extremes — most people "
    "are 60–100). That converts to 0.7 to 4.0 cycles per second (Hz).")
add_para(doc,
    "Anything happening slower than 0.7 Hz (camera slowly drifting, room "
    "lighting changing) gets blocked. Anything happening faster than 4 Hz "
    "(electrical noise, camera jitter) also gets blocked. What's left is "
    "the pulse signal in a clean, isolated frequency band.")

add_heading(doc, "5.8  FFT — Fast Fourier Transform", level=2)
add_para(doc,
    "The FFT is one of the most important algorithms in all of computer science. "
    "It takes a time-series of numbers — a wiggling signal that varies over time — "
    "and decomposes it into the individual frequencies that make it up.")
add_para(doc,
    "We apply the FFT to our cleaned pulse waveform. The output is a graph that "
    "looks like a histogram, with one tall spike. The position of that spike on the "
    "frequency axis is your heart rate. If the spike is at 1.2 Hz, your heart is "
    "beating 1.2 times per second, which means 72 beats per minute.")
add_callout(doc, "Why we trust this",
    "The FFT is the same math used in everything from radio tuners to MP3 "
    "compression to seismology. It's not a heuristic — it's a deterministic "
    "transformation. If the pulse signal is there, the FFT will find it.")

add_heading(doc, "5.9  Peak detection and IBI", level=2)
add_para(doc,
    "Once we have the cleaned pulse waveform, we can do something different "
    "from the FFT: count the individual heartbeats. We use scipy's find_peaks "
    "function to identify each spike in the waveform — each spike is one heartbeat.")
add_para(doc,
    "The gaps between consecutive peaks are called inter-beat intervals (IBIs), "
    "measured in milliseconds. If your peaks are 800 ms apart, your beats-per-minute "
    "is 60000 / 800 = 75 BPM.")

add_heading(doc, "5.10  SDNN — what HRV actually means", level=2)
add_para(doc,
    "Your heart doesn't beat at a perfectly constant rate. Even at rest, the gaps "
    "between beats vary slightly — some are 790 ms, some are 815 ms, some are 805 ms. "
    "This variability is healthy. A heart that beats with too-perfect regularity is "
    "actually a sign of poor autonomic regulation.")
add_para(doc,
    "SDNN stands for \"Standard Deviation of Normal-to-Normal intervals.\" It's "
    "just the statistical standard deviation of all those IBI measurements. A "
    "healthy adult at rest typically has SDNN of 50 milliseconds or more. Below "
    "30 ms is associated with stress, fatigue, or illness.")

add_heading(doc, "5.11  LF/HF ratio and the stress index", level=2)
add_para(doc,
    "If you take the IBI series and do an FFT on it (not the pulse waveform — "
    "the gaps between beats), you can see two important frequency bands:")
add_bullet(doc, "Low Frequency (LF) — 0.04 to 0.15 Hz — driven by a mix of your fight-or-flight system and your rest-and-digest system")
add_bullet(doc, "High Frequency (HF) — 0.15 to 0.40 Hz — driven purely by your rest-and-digest (parasympathetic) system")
add_para(doc,
    "When you're stressed, the fight-or-flight side dominates, which increases LF "
    "relative to HF. So the ratio LF/HF is a real, peer-reviewed proxy for "
    "sympathetic dominance — what we call stress.")
add_para(doc,
    "We take the log of this ratio, push it through a sigmoid function (which "
    "squashes any number into the range 0 to 1), and that's our stress index. "
    "0 means very relaxed; 1 means very stressed.")

doc.add_page_break()

# ============================================================
# 6. THE MATH IN ONE PAGE
# ============================================================
add_heading(doc, "6. The math, all in one place", level=1)

add_para(doc,
    "Here's every formula in the pipeline, with what it does:")

add_heading(doc, "6.1  Mean RGB per frame", level=2)
add_code(doc, "R_t = mean of red values inside ROI at frame t\n"
              "G_t = mean of green values inside ROI at frame t\n"
              "B_t = mean of blue values inside ROI at frame t")

add_heading(doc, "6.2  POS projection", level=2)
add_para(doc, "We normalize each channel by its mean (to remove the DC offset), then apply this fixed matrix:")
add_code(doc, "H = [[ 0,  1, -1],\n"
              "     [-2,  1,  1]]\n\n"
              "S1, S2 = H · normalized(R, G, B)\n"
              "pulse = S1 + (std(S1) / std(S2)) · S2")

add_heading(doc, "6.3  Bandpass filter", level=2)
add_code(doc, "b, a = scipy.signal.butter(order=3, [0.7, 4.0] Hz, btype='band')\n"
              "filtered = scipy.signal.filtfilt(b, a, pulse)")

add_heading(doc, "6.4  Heart rate from FFT", level=2)
add_code(doc, "spectrum = fft(filtered, n=4 * len(filtered))\n"
              "peak_Hz = argmax(spectrum within [60/60, 210/60])\n"
              "heart_rate = peak_Hz · 60")

add_heading(doc, "6.5  SDNN", level=2)
add_code(doc, "peaks = scipy.signal.find_peaks(filtered, prominence=0.3·std(filtered))\n"
              "IBIs = diff(peaks) · 1000 / fps      # in milliseconds\n"
              "SDNN = std(IBIs)")

add_heading(doc, "6.6  Stress index", level=2)
add_code(doc, "psd = welch(resampled IBIs at 4 Hz)\n"
              "LF  = integrate(psd, 0.04..0.15 Hz)\n"
              "HF  = integrate(psd, 0.15..0.40 Hz)\n"
              "stress = sigmoid((log(LF/HF) - 1.0) / scale)")

# ============================================================
# 7. END-TO-END EXAMPLE
# ============================================================
add_heading(doc, "7. A concrete end-to-end example", level=1)

add_para(doc,
    "Let's trace one real scan. Daray clicks Scan live and lets the camera record "
    "for 30 seconds. Here's what happens to the data at each step:")

add_table(doc, [
    ["Stage", "Input", "Output", "Time"],
    ["Browser captures", "30 s of webcam stream", "1 WebM file, ~8 MB", "30 s"],
    ["Upload", "WebM blob", "Multipart HTTP POST", "~2 s"],
    ["FastAPI validates", "Multipart body", "Saved temp file", "<50 ms"],
    ["OpenCV reads frames", "WebM file", "~900 numpy arrays (each 640×480×3)", "~1 s"],
    ["MediaPipe per frame", "Each frame", "468 landmark coordinates", "~2 s total"],
    ["Build ROI polygons", "Landmarks", "Forehead + cheek masks", "<10 ms total"],
    ["Mean RGB", "ROI pixels each frame", "(R, G, B) triple per frame", "<50 ms total"],
    ["Time-series matrix", "900 triples", "(900, 3) numpy array", "<10 ms"],
    ["POS projection", "(900, 3) RGB matrix", "Cleaned pulse waveform", "<50 ms"],
    ["Bandpass filter", "Pulse waveform", "Filtered waveform", "<10 ms"],
    ["FFT", "Filtered waveform", "Heart rate = 72 BPM", "<10 ms"],
    ["Peak detection", "Filtered waveform", "IBIs in milliseconds", "<10 ms"],
    ["SDNN", "IBIs", "HRV = 45 ms", "<5 ms"],
    ["LF/HF stress", "IBIs", "Stress = 0.6", "~20 ms"],
    ["JSON assembly", "All biomarkers", "Shared contract dict", "<5 ms"],
    ["Persist to history", "JSON + scan_id", "Appended to file", "<20 ms"],
    ["Response", "JSON", "HTTP 200 to browser", "<200 ms"],
])

add_para(doc,
    "Total time from \"click Scan live\" to \"see your biomarkers on screen\" is "
    "about 33 to 38 seconds — of which 30 are just the capture itself. The actual "
    "processing on the backend is well under 5 seconds. The site feels fast because "
    "the heavy lifting is fast.")

# ============================================================
# 8. WHY EACH DESIGN CHOICE
# ============================================================
add_heading(doc, "8. Why each design choice was made", level=1)

add_para(doc, "Quick reference for any \"why did you do it this way\" question.")

add_table(doc, [
    ["Choice", "Why"],
    ["MediaPipe FaceMesh (not OpenCV's Haar cascades or dlib)",
     "Real-time on CPU, no GPU needed, more accurate landmarks than older alternatives."],
    ["Forehead + cheeks (not the whole face)",
     "Highest capillary density (strongest pulse signal) and lowest motion artifact."],
    ["POS algorithm (not CHROM or simple green-channel averaging)",
     "Beats CHROM on synthetic data, ties on real data, robust against motion and lighting."],
    ["Butterworth bandpass (not a moving average)",
     "Sharp roll-off, zero phase distortion with filtfilt, mathematically well-behaved."],
    ["FFT for heart rate (not just peak counting)",
     "Robust to missed beats; gives the dominant frequency even with noisy peaks."],
    ["Tightened 60-210 BPM window (not POS paper's 42-240)",
     "Empirically suppresses synthetic-data artifacts; clinical normal is 60-100 anyway."],
    ["Mock blood pressure (not derived)",
     "Classical rPPG physically cannot derive BP without a per-subject cuff calibration."],
    ["FastAPI (not Flask or Django)",
     "Native async, automatic OpenAPI docs, type validation from Python type hints."],
    ["Cloudflare Tunnel (not port-forwarding)",
     "TLS at the edge, DDoS protection, no public IP, firewall stays closed."],
    ["Persistent scan_history.json (not Redis or Postgres)",
     "Right tool for the scale — classroom traffic, ~14 KB max, zero ops overhead."],
])

doc.add_page_break()

# ============================================================
# 9. CONFIDENCE SCORING (per biomarker)
# ============================================================
add_heading(doc, "9. Confidence per biomarker — how much should you trust each number?",
            level=1)

add_para(doc,
    "On 2026-05-25 the professor asked us to return a confidence score alongside "
    "every biomarker, not just the value. Group 3 (risk analysis) wants to weight "
    "their recommendation by how trustworthy the upstream measurement is — a "
    "noisy scan from poor lighting should not drive the same risk call as a clean "
    "scan in good conditions. The contract changed accordingly: every biomarker "
    "is now an object with `value`, `confidence` (0–1), and `unit` keys.")

add_heading(doc, "9.1  Heart rate confidence — FFT peak prominence", level=2)
add_para(doc,
    "The FFT of the cleaned pulse waveform has a peak at the heart-rate frequency. "
    "How big that peak is compared to the average magnitude inside the 60–210 BPM "
    "band tells us whether we found a real periodic signal or just the loudest blob "
    "of noise. The ratio peak_magnitude / mean_magnitude_in_band is the SNR for "
    "this measurement. A clean recording gives a sharp dominant peak (ratio of 5 or "
    "higher); a noisy or motion-corrupted recording produces a peak only marginally "
    "above the noise floor. We map the ratio linearly: ratio = 1.5 → confidence = "
    "0.0 (no signal); ratio = 5.5+ → confidence = 1.0 (very confident).")
add_code(doc, "confidence = clip((peak / mean - 1.5) / 4.0, 0.0, 1.0)")

add_heading(doc, "9.2  HRV and stress confidence — beat count + IBI plausibility",
            level=2)
add_para(doc,
    "SDNN is a sample standard deviation: it gets more reliable the more beats we "
    "detect. Over 30 seconds at a normal heart rate we expect 30–40 beats; we treat "
    "20 detected beats as full confidence for HRV and 25 for stress (the LF/HF "
    "spectral analysis needs more samples than the time-domain SDNN).")
add_para(doc,
    "On top of beat count we add a plausibility check: the coefficient of variation "
    "of the inter-beat-interval series. Real human HRV produces a CV in the 0.05–"
    "0.15 range. If the CV spikes above 0.30 we are probably missing or extra-"
    "counting beats, so we apply a linear penalty (floored at 0.2 so we never go "
    "all the way to zero unless we have no beats at all).")
add_code(doc, "hrv_confidence  = min(1, n_peaks / 20) * ibi_plausibility_penalty\n"
              "stress_confidence = min(1, n_peaks / 25) * ibi_plausibility_penalty")

add_heading(doc, "9.3  Blood pressure confidence — a low floor by design", level=2)
add_para(doc,
    "We return a blood-pressure value because the contract requires it, but "
    "classical rPPG physically cannot measure BP without a per-subject cuff "
    "calibration reference. We report a fixed confidence of 0.10 — low enough that "
    "downstream consumers know not to trust it, but non-zero so it can still be "
    "rendered. The response also carries a `note` field that explains the "
    "limitation in plain English.")

add_heading(doc, "9.4  How Group 3 should consume confidence", level=2)
add_para(doc,
    "Treat confidence as a multiplicative weight on the biomarker's contribution "
    "to risk scoring. A heart rate of 92 BPM at confidence 0.9 should drive risk "
    "more strongly than the same value at confidence 0.4. A flat threshold "
    "(\"high risk if HR > 90\") that ignores confidence will produce false alarms "
    "from noisy scans.")

doc.add_page_break()

# ============================================================
# 10. COMMON QUESTIONS — what the professor asked on 2026-05-25
# ============================================================
add_heading(doc, "10. Common questions, answered", level=1)

add_para(doc,
    "These are the questions that came up in the 2026-05-25 progress review. We "
    "answer each one explicitly here so anyone reading the project later can find "
    "the answer in one place rather than scrubbing the meeting recording.")

add_heading(doc, "10.1  What is RMSE, and why use MAE and RMSE together?", level=2)
add_para(doc,
    "MAE (mean absolute error) is the average of |predicted − actual| over every "
    "subject in the test set. RMSE (root mean square error) is the square root of "
    "the average of (predicted − actual)². They tell you slightly different things:")
add_bullet(doc, "MAE = how far off we are, on average. Easy to interpret — "
                "\"the pipeline is off by about 4 BPM on a typical subject.\"")
add_bullet(doc, "RMSE = how far off we are when we are wrong, weighted by how badly "
                "we are wrong. Because the errors are squared before averaging, a "
                "single 30-BPM miss hurts RMSE much more than ten 3-BPM misses do. "
                "RMSE is always greater than or equal to MAE; the gap between them "
                "tells you whether your errors are uniformly distributed (small gap) "
                "or have outliers (large gap).")
add_para(doc,
    "On our UBFC evaluation, POS reaches MAE 4.06 BPM and RMSE 7.78 BPM. The 1.9× "
    "gap reflects the four outlier subjects we documented in the writeup. Both "
    "metrics are reported because each answers a different question.")

add_heading(doc, "10.2  Train / test / evaluation split", level=2)
add_para(doc,
    "For a classical signal-processing pipeline like POS, the split conversation "
    "is different from a deep-learning model. POS has no learned parameters — "
    "the only \"hyperparameter\" is the empirically tightened 60–210 BPM FFT "
    "search window, which we set on a small validation subset of SCAMPS to "
    "suppress synthetic-animation artifacts. There is no training step, so there "
    "is no train/test/eval three-way split in the usual ML sense.")
add_para(doc,
    "What we do report is held-out evaluation on the full UBFC-rPPG real-human "
    "dataset (n = 42 subjects). Because no model was trained on those subjects, "
    "the entire UBFC set functions as the held-out test set. If we later add a "
    "deep-learning extension (the stretch goal — see Task 5 of the project "
    "brief), the split will become an explicit 70/15/15 train/val/test on the "
    "SCAMPS training corpus, with UBFC reserved for cross-domain evaluation.")

add_heading(doc, "10.3  Why a mathematical approach instead of training a model?",
            level=2)
add_para(doc,
    "We considered both. The classical mathematical approach (POS algorithm + FFT "
    "+ peak detection) was chosen as the primary method for four reasons:")
add_bullet(doc, "Training data — published deep-learning rPPG models are trained "
                "on hundreds of subject-hours. The only large open dataset we can "
                "access without institutional negotiation is SCAMPS (synthetic), "
                "which transfers poorly to real subjects. UBFC has 42 real "
                "subjects — enough to evaluate, not to train from scratch.")
add_bullet(doc, "Interpretability — POS produces a human-readable intermediate "
                "pulse waveform we can inspect frame-by-frame. A trained model "
                "emits an opaque waveform that is hard to debug when accuracy "
                "regresses.")
add_bullet(doc, "Course learning outcomes — the rubric (CLO 3 & 4) asks for "
                "demonstrated understanding of signal-processing fundamentals. "
                "Implementing POS, Butterworth filtering, FFT, and peak detection "
                "by hand makes those fundamentals legible. Calling "
                "model.forward() does not.")
add_bullet(doc, "Real-time on CPU — POS runs in well under a second on a 30-s "
                "video on CPU. Deep models typically need a GPU.")
add_para(doc,
    "We have scaffolded the deep-learning comparison (PhysNet / TS-CAN via the "
    "rPPG-Toolbox, Liu et al. 2023) as the stretch goal. The classical pipeline "
    "already clears the rubric's MAE < 10 BPM target at 4.06 BPM, so the deep "
    "comparison is additive rather than load-bearing.")

add_heading(doc, "10.4  How does the dataset map videos to biometric values?",
            level=2)
add_para(doc,
    "Each subject in the UBFC-rPPG dataset has two files: a video file and a "
    "ground_truth.txt file. The text file contains three rows: row 1 is the raw "
    "PPG waveform from a synchronized finger pulse-oximeter, row 2 is the "
    "instantaneous heart rate the oximeter device reported (sample-by-sample), "
    "and row 3 is the timestamp in scientific notation. Every video frame at time "
    "t has a matched ground-truth HR value at the same t — that is what lets us "
    "compute MAE per subject. SCAMPS has the same structure but the ground truth "
    "is simulated rather than measured.")

add_heading(doc, "10.5  Why MediaPipe FaceMesh instead of dlib or MTCNN?", level=2)
add_para(doc,
    "Two reasons. (1) Granularity: FaceMesh emits 468 dense landmarks per frame, "
    "letting us mask the forehead and both cheeks precisely. dlib (68 landmarks) "
    "and MTCNN (5 landmarks) only return bounding boxes or sparse points — "
    "neither is enough to draw a polygon over capillary-rich skin. (2) Real-time "
    "on CPU: FaceMesh runs at video framerate on a regular laptop. RetinaFace is "
    "more accurate at extreme angles but typically needs a GPU and is overkill "
    "for our use case (cooperative subject, mostly frontal).")

doc.add_page_break()

# ============================================================
# 11. ONE PARAGRAPH SUMMARY
# ============================================================
add_heading(doc, "11. The whole project, in one paragraph", level=1)

add_para(doc,
    "You open the site, allow camera access, and let the browser record 30 seconds "
    "of you sitting still. That WebM file is uploaded over an HTTPS connection that "
    "passes through Cloudflare's edge and into a tunnel back to my homelab. The "
    "FastAPI backend saves the upload, opens it with OpenCV, and walks frame by "
    "frame through the video. Each frame is fed to MediaPipe's FaceMesh model, "
    "which returns 468 landmark coordinates marking specific anatomical points on "
    "your face. We connect specific landmarks into polygons over your forehead and "
    "both cheeks — the regions where blood-flow color changes are strongest — and "
    "we compute the average red, green, and blue values inside those regions. After "
    "900 frames we have a 900-by-3 table of RGB averages. The POS algorithm projects "
    "this table to mathematically separate the pulse signal from lighting and motion "
    "noise, a Butterworth bandpass filter removes anything outside the physiological "
    "heart-rate range, and an FFT finds the dominant frequency — which, multiplied "
    "by 60, is your heart rate in BPM. The same cleaned waveform is peak-detected "
    "to find individual beats; the gaps between beats give us HRV (SDNN), and the "
    "low-frequency over high-frequency ratio of those gaps gives us a stress score. "
    "We package the four biomarkers into the shared JSON contract, append them to "
    "scan_history.json with a generated scan ID and your name, and send the JSON "
    "back to your browser, which renders the dashboard. Groups 3 and 4 can later "
    "retrieve your scan or anyone else's via GET /api/biometrics, hand it to their "
    "models, and tell you whether to eat the soup.")

# --- Save ---
doc.save(OUT)
print(f"wrote {OUT}")
