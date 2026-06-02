"""Update Daray's Week 4 docx by inserting architecture diagrams in the right
positions, then export to .docx, .md, and .pdf alongside it.

Input:    ~/Downloads/Week4_Project_Model_Development_Framework_Design.docx
Outputs:  docs/Week4_Project_Model_Development_Framework_Design.docx
          docs/Week4_Project_Model_Development_Framework_Design.md
          docs/Week4_Project_Model_Development_Framework_Design.pdf

The two images inserted are the same SVGs we ship in docs/figures/:
  figure1-signal-chain.png  → "Through the Pipeline" data-flow section
  figure2-architecture.png  → "Pipeline Architecture Overview" section
"""
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = Path.home() / "Downloads" / "Week4_Project_Model_Development_Framework_Design.docx"
DOCX_OUT = ROOT / "docs" / "Week4_Project_Model_Development_Framework_Design.docx"
MD_OUT = DOCX_OUT.with_suffix(".md")
PDF_OUT = DOCX_OUT.with_suffix(".pdf")
MEDIA_OUT = ROOT / "docs" / "wk4_media"

FIG_SIGNAL = ROOT / "docs" / "figures" / "figure1-signal-chain.png"
FIG_ARCH = ROOT / "docs" / "figures" / "figure2-architecture.png"


def insert_picture_after(doc, target_para, image_path, caption_text,
                         width_inches=6.5):
    """Insert a picture + centered italic caption right after target_para."""
    # Create new image paragraph at end of doc, then relocate via XML.
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption_text)
    cap_run.italic = True
    cap_run.font.size = Pt(10)

    # Relocate the two new paragraphs immediately after target_para.
    body = doc.element.body
    body.remove(img_p._element)
    body.remove(cap_p._element)
    target_para._element.addnext(cap_p._element)
    target_para._element.addnext(img_p._element)


def find_paragraph_starting_with(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def find_paragraph_containing(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def fix_professor_name(doc):
    """Daray's draft listed the professor as 'Admiluyi'; the correct name is
    'Desmond Ademiluyi'."""
    fixed = 0
    for p in doc.paragraphs:
        new_text = p.text.replace("Admiluyi", "Desmond Ademiluyi")
        if new_text != p.text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            fixed += 1
    if fixed:
        print(f"  professor name corrected in {fixed} paragraph(s)")


def update_docx():
    print(f"Reading {SRC}")
    doc = Document(SRC)
    fix_professor_name(doc)

    # 1. Pipeline Architecture Overview → architecture/system diagram
    arch_anchor = find_paragraph_starting_with(doc, "Five-Stage Design")
    if arch_anchor is None:
        raise SystemExit("Couldn't find 'Five-Stage Design' anchor")

    # 2. Through the Pipeline (Data Flow section) → signal-chain diagram
    flow_anchor = find_paragraph_containing(
        doc, "A visual diagram will be included"
    )
    if flow_anchor is None:
        # Fallback: insert at the 'Through the Pipeline' header
        flow_anchor = find_paragraph_starting_with(doc, "Through the Pipeline")
    if flow_anchor is None:
        raise SystemExit("Couldn't find Data Flow anchor")

    # Replace Daray's "visual diagram will be included" placeholder with the
    # actual signal-chain figure + caption.
    flow_replacement_text = (
        "The signal-chain diagram in Figure 1 below visualizes the data flow "
        "through stages 1–4. Tensor shapes are annotated at each transition "
        "so the per-stage transform is traceable end-to-end."
    )
    for run in flow_anchor.runs:
        run.text = ""
    flow_anchor.runs[0].text = flow_replacement_text if flow_anchor.runs else ""
    if not flow_anchor.runs:
        flow_anchor.add_run(flow_replacement_text)

    # Insert signal-chain figure after the rewritten "Through the Pipeline" line
    insert_picture_after(
        doc, flow_anchor, FIG_SIGNAL,
        caption_text="Figure 1. rPPG signal-chain — data flow through stages 1–4.",
        width_inches=6.5,
    )

    # Insert architecture figure after the "Five-Stage Design" intro.
    # Use ASCII arrows in the caption so the LaTeX PDF engine doesn't choke
    # on missing Unicode glyphs in italic Helvetica.
    insert_picture_after(
        doc, arch_anchor, FIG_ARCH,
        caption_text="Figure 2. System architecture - client to Cloudflare tunnel "
                     "to containerized backend to JSON contract for Groups 3 & 4.",
        width_inches=6.5,
    )

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


def export_markdown():
    """pandoc DOCX → Markdown, with media extracted to a sibling folder."""
    if MEDIA_OUT.exists():
        shutil.rmtree(MEDIA_OUT)
    cmd = [
        "pandoc", str(DOCX_OUT),
        "-f", "docx", "-t", "gfm",
        f"--extract-media={MEDIA_OUT}",
        "-o", str(MD_OUT),
    ]
    subprocess.run(cmd, check=True)
    print(f"Wrote {MD_OUT}  (media → {MEDIA_OUT.name}/)")


def export_pdf():
    """pandoc DOCX → PDF via xelatex."""
    cmd = [
        "pandoc", str(DOCX_OUT),
        "--pdf-engine=xelatex",
        "-V", "geometry:margin=1in",
        "-V", "mainfont=Helvetica",
        "-V", "monofont=Menlo",
        "-V", "fontsize=11pt",
        "-V", "linkcolor=blue",
        "-o", str(PDF_OUT),
    ]
    subprocess.run(cmd, check=True)
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    update_docx()
    export_markdown()
    export_pdf()
